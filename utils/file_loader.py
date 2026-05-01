"""
File loading utilities.

Loads resume from multiple formats: PDF, DOCX, and TXT.
Handles format detection and text extraction.
"""

from typing import Optional
import fitz  # PyMuPDF
from docx import Document

from utils.logger import get_logger

logger = get_logger(__name__)


def load_pdf(file_path: str) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
        
    Raises:
        RuntimeError: If PDF reading fails
    """
    try:
        logger.info(f"Loading PDF: {file_path}")
        text = ""
        
        doc = fitz.open(file_path)
        page_count = doc.page_count
        logger.debug(f"PDF has {page_count} pages")
        
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            text += page_text
            logger.debug(f"Extracted page {page_num}")
        
        doc.close()
        
        extracted_length = len(text.strip())
        logger.info(f"PDF extraction complete: {extracted_length} characters")
        return text.strip()
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"Error reading PDF file: {e}")


def load_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        RuntimeError: If DOCX reading fails
    """
    try:
        logger.info(f"Loading DOCX: {file_path}")
        doc = Document(file_path)
        text = []
        
        paragraph_count = len(doc.paragraphs)
        logger.debug(f"DOCX has {paragraph_count} paragraphs")
        
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        
        result = "\n".join(text)
        extracted_length = len(result.strip())
        logger.info(f"DOCX extraction complete: {extracted_length} characters")
        return result
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise RuntimeError(f"Error reading DOCX file: {e}")


def load_text_file(file_path: str) -> str:
    """
    Extract text from TXT file.
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        File content
        
    Raises:
        RuntimeError: If file reading fails
    """
    try:
        logger.info(f"Loading TXT: {file_path}")
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read().strip()
        
        extracted_length = len(content)
        logger.info(f"TXT extraction complete: {extracted_length} characters")
        return content
        
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        raise RuntimeError(f"Error reading TXT file: {e}")


def load_resume_file(file_path: str) -> str:
    """
    Load resume from PDF, DOCX, or TXT format.
    
    Automatically detects file format by extension.
    
    Args:
        file_path: Path to resume file
        
    Returns:
        Extracted resume text
        
    Raises:
        ValueError: If file format is unsupported
        RuntimeError: If file reading fails
    """
    file_path_lower = file_path.lower()
    
    logger.info(f"Loading resume: {file_path}")
    
    try:
        if file_path_lower.endswith(".pdf"):
            return load_pdf(file_path)
        elif file_path_lower.endswith(".docx"):
            return load_docx(file_path)
        elif file_path_lower.endswith(".txt"):
            return load_text_file(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {file_path}. "
                "Use PDF, DOCX, or TXT."
            )
    except Exception as e:
        logger.error(f"Resume loading failed: {e}")
        raise
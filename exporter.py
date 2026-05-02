"""Safe export helpers for TXT, DOCX, and PDF."""

from __future__ import annotations

import textwrap
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Anchor output directory relative to this file, not the process cwd
OUTPUT_DIR = Path(__file__).parent / "outputs"

# Heading detection: all-caps or title-case short lines are treated as section headers
_HEADING_PATTERN = lambda line: (
    len(line) <= 40
    and (line.isupper() or (line.istitle() and len(line.split()) <= 5))
)


def export_resume(text: str, output_format: str = "docx") -> str:
    output_format = output_format.lower().strip()
    if output_format not in {"docx", "pdf", "txt"}:
        raise ValueError("Export format must be docx, pdf, or txt.")
    if not text.strip():
        raise ValueError("Cannot export an empty resume.")

    OUTPUT_DIR.mkdir(exist_ok=True)
    stem = f"optimized_resume_{uuid.uuid4().hex[:8]}"
    path = OUTPUT_DIR / f"{stem}.{output_format}"

    if output_format == "docx":
        return _export_docx(text, path)
    if output_format == "pdf":
        return _export_pdf(text, path)

    # TXT — plain passthrough
    path.write_text(text, encoding="utf-8")
    return str(path)


def _export_docx(text: str, path: Path) -> str:
    doc = Document()

    # Set default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    prev_blank = False
    for raw_line in text.splitlines():
        line = raw_line.strip()

        if not line:
            # Collapse consecutive blank lines; insert spacing via paragraph spacing instead
            if not prev_blank:
                doc.add_paragraph()
            prev_blank = True
            continue

        prev_blank = False

        if _HEADING_PATTERN(line):
            heading = doc.add_paragraph()
            run = heading.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
        else:
            para = doc.add_paragraph()
            para.add_run(line)

    doc.save(path)
    return str(path)


def _export_pdf(text: str, path: Path) -> str:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    _, height = letter
    x = 50
    y = height - 60
    line_height = 15
    heading_height = 17

    pdf.setFont("Helvetica", 11)

    for raw_line in text.splitlines():
        line = raw_line.strip()

        if not line:
            y -= line_height // 2   # half-line spacing for blank lines
            continue

        is_heading = _HEADING_PATTERN(line)
        wrapped = textwrap.wrap(line, width=90)

        for i, segment in enumerate(wrapped or [line]):
            if y < 60:
                pdf.showPage()
                pdf.setFont("Helvetica", 11)
                y = height - 60

            if is_heading and i == 0:
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawString(x, y, segment)
                pdf.setFont("Helvetica", 11)
                y -= heading_height
            else:
                pdf.drawString(x, y, segment)
                y -= line_height

    pdf.save()
    return str(path)